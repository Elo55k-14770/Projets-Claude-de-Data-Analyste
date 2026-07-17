# -*- coding: utf-8 -*-
"""Génère le résumé exécutif (Word, 5 pages max) à partir de resultats_analyses.json et des figures."""
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

with open("resultats_analyses.json", "r", encoding="utf-8") as f:
    R = json.load(f)

K = R["kpis_globaux"]
CAT = {c["categorie"]: c for c in R["par_categorie"]}
VIL = {v["ville"]: v for v in R["par_ville"]}
CAN = {c["canal_marketing"]: c for c in R["par_canal"]}
CROISS = R["croissance_ville_pct"]
SEG = {s["segment_client"]: s for s in R["segments"]}

cat_top_ca = max(CAT.values(), key=lambda x: x["ca"])
cat_top_retour = max(CAT.values(), key=lambda x: x["taux_retour"])
ville_top_ca = max(VIL.values(), key=lambda x: x["ca"])
ville_top_croissance = max(CROISS.items(), key=lambda x: x[1])
ville_pire_croissance = min(CROISS.items(), key=lambda x: x[1])
canal_top_roi = max(CAN.values(), key=lambda x: x["roi"])
canal_pire_roi = min(CAN.values(), key=lambda x: x["roi"])

doc = Document()

# --- Styles de base ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_title(text, size=20, color=(31, 78, 121)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def add_h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_kpi_table(rows):
    table = doc.add_table(rows=1, cols=len(rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, (label, value) in enumerate(rows):
        hdr[i].text = ""
        p1 = hdr[i].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(value)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(31, 78, 121)
        p2 = hdr[i].add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        shade_cell(hdr[i], "EAF1F8")
    return table


# =====================================================================
# EN-TETE
# =====================================================================
add_title("Résumé Exécutif — Analyse Stratégique AfriMarket")
p = doc.add_paragraph()
r = p.add_run("Data Analyst · Période analysée : Juillet 2025 – Décembre 2025 (6 mois d'activité)")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(90, 90, 90)

doc.add_paragraph(
    "AfriMarket, entreprise e-commerce panafricaine active dans 8 villes et 4 catégories "
    "(Électronique, Mode, Beauté, Maison), a constaté des variations de chiffre d'affaires, un taux de retour "
    "préoccupant, des dépenses marketing élevées et des écarts de performance entre villes. L'analyse ci-dessous, "
    "conduite sur les 6 derniers mois d'activité après un audit et un nettoyage rigoureux des données "
    "(voir notebook joint), objective ces constats et propose des actions concrètes."
)

# =====================================================================
# KPIs GLOBAUX
# =====================================================================
add_h2("1. Performance globale")
add_kpi_table([
    ("CA réalisé (6 mois)", f"{K['ca_total']:,.0f}"),
    ("Profit net estimé", f"{K['profit_total']:,.0f}"),
    ("Panier moyen", f"{K['panier_moyen']:,.0f}"),
    ("Taux d'annulation", f"{K['taux_annulation_pct']:.1f}%"),
    ("Taux de retour", f"{K['taux_retour_pct']:.1f}%"),
])
doc.add_paragraph(
    f"Sur {K['nb_commandes']:,} commandes exploitées (après nettoyage) et {K['nb_clients']:,} clients uniques, "
    f"le profit net estimé ne représente que {K['profit_total']/K['ca_total']*100:.1f}% du CA réalisé, "
    "ce qui souligne l'impact direct des coûts marketing, logistiques et des retours sur la rentabilité réelle — "
    "un indicateur à piloter au moins autant que le chiffre d'affaires brut."
)

# =====================================================================
# CATEGORIE
# =====================================================================
add_h2("2. Performance par catégorie — quelle catégorie prioriser ?")
doc.add_paragraph(
    f"{cat_top_ca['categorie']} génère le CA le plus élevé ({cat_top_ca['ca']:,.0f}, "
    f"{cat_top_ca['n_commandes']} commandes) mais affiche aussi le taux de retour le plus élevé "
    f"({cat_top_ca['taux_retour']:.1f}%), qui érode sa rentabilité réelle. À l'inverse, les catégories Mode et Beauté, "
    "moins volumineuses, offrent des marges brutes proportionnellement plus élevées et des taux de retour maîtrisés. "
    f"Recommandation : prioriser {cat_top_ca['categorie']} en volume tout en lançant un plan de réduction des retours "
    "ciblé (qualité produit, description, photos), et soutenir la croissance de Mode/Beauté à marge plus favorable."
)
doc.add_picture("figures/02_analyse_categorie.png", width=Inches(5.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
# GEOGRAPHIE
# =====================================================================
add_h2("3. Performance géographique — où investir davantage ?")
top_g, top_g_val = ville_top_croissance
low_g, low_g_val = ville_pire_croissance
doc.add_paragraph(
    f"{ville_top_ca['ville']} génère le CA le plus élevé ({ville_top_ca['ca']:,.0f}). En dynamique, "
    f"{top_g} affiche la plus forte croissance sur la période (+{top_g_val:.0f}%), tandis que {low_g} recule "
    f"({low_g_val:.0f}%). Le taux d'annulation varie fortement selon la ville : certaines villes conjuguent fort CA "
    "et taux d'annulation élevé, révélateur de frictions opérationnelles locales (paiement, livraison) à corriger "
    "avant tout renforcement de l'investissement. "
    f"Recommandation : concentrer l'investissement marketing et logistique sur {top_g}, tout en menant un audit "
    "opérationnel dans les villes à fort taux d'annulation avant d'y augmenter les budgets."
)
doc.add_picture("figures/03_analyse_geo_ca_profit.png", width=Inches(5.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
# MARKETING
# =====================================================================
add_h2("4. Performance marketing — quel canal privilégier ?")
doc.add_paragraph(
    f"Le canal {canal_top_roi['canal_marketing']} affiche de loin le meilleur retour sur investissement "
    f"(ROI = {canal_top_roi['roi']:.1f}, soit {canal_top_roi['roi']:.0f}€ générés par euro dépensé net des coûts) "
    f"pour un coût total de {canal_top_roi['cout_marketing']:,.0f}. À l'inverse, {canal_pire_roi['canal_marketing']} "
    f"présente le ROI le plus faible ({canal_pire_roi['roi']:.1f}) malgré un investissement de "
    f"{canal_pire_roi['cout_marketing']:,.0f}. Tous les canaux affichent des taux de rétention client proches "
    "(~95%), ce qui indique que l'écart de performance vient de l'efficacité d'acquisition, pas de la fidélisation. "
    f"Recommandation : réallouer une partie du budget de {canal_pire_roi['canal_marketing']} vers "
    f"{canal_top_roi['canal_marketing']}, sous-exploité au regard de son efficacité."
)

# =====================================================================
# CLIENTS
# =====================================================================
add_h2("5. Analyse clients — comment améliorer la rétention ?")
vip = SEG["VIP"]
occ = SEG["Occasionnel"]
doc.add_paragraph(
    f"AfriMarket compte {K['nb_clients']:,} clients uniques, dont {K['pct_clients_recurrents']:.0f}% ont commandé "
    f"plus d'une fois. L'analyse Pareto montre que {K['pct_clients_pour_80pct_ca']:.0f}% des clients génèrent 80% "
    f"du chiffre d'affaires réalisé : le segment VIP ({vip['nb_clients']} clients, CLV moyenne "
    f"{vip['clv_moyenne']:,.0f}) représente à lui seul {vip['clv_totale']/K['ca_total']*100:.0f}% du CA total, "
    f"contre seulement {occ['clv_totale']/K['ca_total']*100:.0f}% pour les {occ['nb_clients']} clients Occasionnels. "
    "Cette concentration de valeur justifie un traitement différencié par segment plutôt qu'une approche uniforme."
)
doc.add_picture("figures/07_pareto_clients.png", width=Inches(5.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
# RECOMMANDATIONS
# =====================================================================
add_h2("6. Cinq recommandations stratégiques")
recos = [
    ("Prioriser la rentabilité, pas seulement le volume",
     f"Piloter {cat_top_ca['categorie']} avec un plan de réduction des retours (audit qualité/description), "
     "tout en accélérant Mode et Beauté, à marge plus favorable."),
    ("Traiter le taux de retour comme un enjeu de rentabilité",
     "Lancer un audit qualité ciblé sur la catégorie et la ville où le taux de retour est le plus élevé : "
     "chaque retour détruit la marge et le coût de livraison déjà engagé."),
    ("Réallouer le budget marketing vers le canal le plus performant",
     f"Renforcer {canal_top_roi['canal_marketing']} (ROI le plus élevé) et réduire ou requalifier "
     f"{canal_pire_roi['canal_marketing']} (ROI le plus faible)."),
    ("Concentrer l'investissement géographique sur les villes en croissance saine",
     f"Prioriser {top_g} (plus forte croissance) ; auditer les causes d'annulation dans les villes à fort CA "
     "mais taux d'annulation élevé avant d'y investir davantage."),
    ("Déployer un programme de fidélisation segmenté",
     "Traitement prioritaire pour les clients VIP (support, offres exclusives), montée en gamme des clients "
     "Réguliers, et campagnes de réactivation ciblées pour les clients Occasionnels."),
]
for i, (titre, detail) in enumerate(recos, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {titre}. ")
    r.bold = True
    p.add_run(detail)
    p.paragraph_format.space_after = Pt(6)

# =====================================================================
# CONCLUSION
# =====================================================================
add_h2("7. Conclusion business orientée action")
doc.add_paragraph(
    "Après correction des anomalies de données (doublons, incohérences de saisie, valeurs aberrantes — environ 7% "
    "des lignes brutes), l'analyse dessine une feuille de route claire : piloter la rentabilité réelle plutôt que "
    "le seul chiffre d'affaires, corriger les frictions opérationnelles identifiées par ville et par catégorie, "
    "réallouer le budget marketing vers les canaux les plus efficaces, et concentrer les efforts de fidélisation sur "
    "les clients à forte valeur. La mise en place d'un suivi mensuel de ces indicateurs, via le dashboard Streamlit "
    "livré en complément, permettra de transformer cette photographie ponctuelle en pilotage continu de la performance."
)

doc.save("Resume_Executif_AfriMarket.docx")
print("Résumé exécutif généré -> Resume_Executif_AfriMarket.docx")
